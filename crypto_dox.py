import pandas as pd
from docx import Document

# Load the Excel file
file_path = "crypto_data.xlsx"
df = pd.read_excel(file_path)

# Create a new Word document
doc = Document()
doc.add_heading("Cryptocurrency Market Analysis Report", level=1)

# Add Top 5 Cryptos by Market Cap
doc.add_heading("Top 5 Cryptocurrencies by Market Cap", level=2)
top_5_market_cap = df.nlargest(5, "Market Cap")[["Symbol", "Market Cap"]]

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Symbol"
hdr_cells[1].text = "Market Cap (USD)"

for _, row in top_5_market_cap.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row["Symbol"]
    row_cells[1].text = f"${row['Market Cap']:,.2f}"

# Add Average Price
doc.add_heading("Average Price of Cryptocurrencies", level=2)
average_price = df["Price (USD)"].mean()
doc.add_paragraph(f"The average price of all cryptocurrencies is: **${average_price:,.2f}**")

# Add Highest & Lowest 24h Price Change
doc.add_heading("24h Price Change Analysis", level=2)

highest_24h_change = df.loc[df["24h Price Change (%)"].idxmax(), ["Symbol", "24h Price Change (%)"]]
lowest_24h_change = df.loc[df["24h Price Change (%)"].idxmin(), ["Symbol", "24h Price Change (%)"]]

doc.add_paragraph(f"📈 **Highest 24h Change:** {highest_24h_change['Symbol']} ({highest_24h_change['24h Price Change (%)']:.2f}%)")
doc.add_paragraph(f"📉 **Lowest 24h Change:** {lowest_24h_change['Symbol']} ({lowest_24h_change['24h Price Change (%)']:.2f}%)")

# Save the Word document
report_path = "Crypto_Report.docx"
doc.save(report_path)

print(f"Report saved as {report_path}")
